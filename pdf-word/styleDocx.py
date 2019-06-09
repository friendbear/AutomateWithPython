#! python3
import docx
import setup

# Python-Docx Style
# see: https://python-docx.readthedocs.io/en/latest/user/styles-using.html

doc = docx.Document('demo.docx')
doc.paragraphs[0].style = 'Normal'
doc.paragraphs[1].runs[0].style = 'Quote Char'
doc.paragraphs[1].runs[1].underline = True
doc.paragraphs[1].runs[2].underline = True

doc.save('restyled.docx')

