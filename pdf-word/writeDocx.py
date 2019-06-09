#! python3

import docx
import setup

doc = docx.Document()

doc.add_paragraph('Hello world!')
doc.add_paragraph('これは第２段落です。').add_run('これは第２段落に追加したテキストです。')
doc.add_paragraph('これは第３段落です。')

# header
doc.add_heading('Header 0', 0)
doc.add_heading('Header 1', 1)
doc.add_heading('Header 2', 2)
doc.add_heading('Header 3', 3)
doc.add_heading('Header 4', 4)

# break page_break
doc.add_page_break()
doc.add_paragraph('page 2')

doc.add_page_break()

# add image
doc.add_picture('zophie.png', width=docx.shared.Inches(1), height=docx.shared.Cm(4))


doc.save('document.docx')
